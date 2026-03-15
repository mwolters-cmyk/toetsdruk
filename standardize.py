#!/usr/bin/env python3
"""
Studiewijzer Standardizer
Leest studiewijzers in diverse formats (docx, pdf, txt, xlsx) en zet ze om
naar gestandaardiseerd JSON met alle toetsmomenten, via de Claude API.
"""

import argparse
import asyncio
import json
import os
import re
import sys
import time
from pathlib import Path

import anthropic
import fitz  # PyMuPDF
import openpyxl
from docx import Document

# ── Configuratie ──────────────────────────────────────────────────────────────

BASE_DIR = Path(__file__).parent
EXTRACTED_DIR = BASE_DIR / "_extracted"
OUTPUT_DIR = BASE_DIR / "output"
MODEL = "claude-haiku-4-5-20251001"
MAX_CONCURRENT = 5
SKIP_EXTENSIONS = {".url", ".kmz", ".pptx", ".doc", ".rtf"}

VAK_AFKORTINGEN = {
    "aardrijkskunde": "ak", "biologie": "biol", "duits": "dutl",
    "economie": "econ", "engels": "entl", "filosofie": "fi",
    "frans": "fatl", "geschiedenis": "ges", "grieks": "gtc",
    "kunst - bv": "kubv", "kunst - dr": "kudr", "latijn": "ltc",
    "maatschappijleer": "maat", "muziek": "mu", "natuurkunde": "nat",
    "nederlands": "netl", "scheikunde": "schk",
    "wiskunde a": "wisa", "wiskunde b": "wisb", "wiskunde c": "wisc",
    "wiskunde d": "wisd", "wiskunde": "wi",
    "drama": "kudr", "beeldende vorming": "kubv",
    "onderzoek": "onderzoek", "bv": "kubv",
}

SYSTEM_PROMPT = """\
Je bent een parser die studiewijzers van een Nederlandse middelbare school (gymnasium) \
omzet naar gestructureerd JSON.

SCHOOLJAAR 2025-2026 structuur:
- Onderbouw (klas 1-3): 3 modules per jaar
  - Module 1: week 36-48, toetsweek week 47-48 (17-26 november 2025)
  - Module 2: week 49-11, toetsweek week 10-11 (4-13 maart 2026)
  - Module 3: week 12-26, toetsweek week 25-26 (juni 2026)
- Bovenbouw klas 4-5: zelfde modules 1-3
- Klas 6: modules 4, 5, 6 (zelfde periodes als 1, 2, 3)

TOETSTYPEN die je moet herkennen:
- proefwerk (pw): schriftelijk examen in de toetsweek, typisch 50-100 min
- so: schriftelijke overhoring, kort, tussendoor
- uso: uitgebreide schriftelijke overhoring, langer dan SO
- po: praktische opdracht (kan meerdere weken duren)
- mondeling: mondelinge toets/overhoring
- portfolio: doorlopende verzameling werk
- handelingsdeel (hd): voldaan/niet-voldaan onderdeel
- presentatie: presentatieopdracht
- oefentoets: diagnostische toets, oefentoets, d-toets, nulmeting, formatieve toets, \
proeftoets, quiz (muziek e.d.), practice test — NIET summatief maar wel relevant om te registreren
- anders: alles wat niet in bovenstaande past

INSTRUCTIES:
1. Identificeer ALLE toetsmomenten in de tekst.
2. Voor elk toetsmoment, extraheer: type, beschrijving, weeknummer, datum, duur, weging, \
of het in de toetsweek valt, en de toetsstof.
3. Geef bij elk toetsmoment een zekerheidsscore:
   - "hoog": expliciet genoemd met week/datum
   - "middel": af te leiden uit context (bijv. positie in weekplanning)
   - "laag": impliciet of onduidelijk
4. Wees CONSERVATIEF: liever een toets missen dan er een verzinnen.
5. Als de tekst geen toetsen bevat (bijv. alleen een begrippenlijst), retourneer een lege lijst.
6. Huiswerk en oefenopgaven (die niet als toets benoemd worden) zijn GEEN toetsen.
7. Diagnostische toetsen, oefentoetsen, d-toetsen, nulmetingen, proeftoetsen en quizzes \
(bijv. "proefwerk-quiz", "pw quiz", "Kerstquiz") WEL opnemen als type "oefentoets" — \
ook als ze niet meetellen voor het cijfer. Let op: als een quiz onderdeel is van een \
presentatie (bijv. "Presentatie + quiz"), classificeer dan als "presentatie", niet "oefentoets".
"""

JSON_SCHEMA = {
    "type": "object",
    "properties": {
        "toetsen": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "type": {
                        "type": "string",
                        "enum": ["proefwerk", "so", "uso", "po", "mondeling",
                                 "portfolio", "handelingsdeel", "presentatie",
                                 "oefentoets", "anders"]
                    },
                    "beschrijving": {"type": "string"},
                    "week": {"type": ["integer", "null"]},
                    "datum": {"type": ["string", "null"]},
                    "duur_min": {"type": ["integer", "null"]},
                    "weging": {"type": ["number", "null"]},
                    "in_toetsweek": {"type": "boolean"},
                    "stof": {"type": ["string", "null"]},
                    "zekerheid": {
                        "type": "string",
                        "enum": ["hoog", "middel", "laag"]
                    },
                },
                "required": ["type", "beschrijving", "week", "in_toetsweek", "zekerheid"],
            },
        },
        "opmerkingen": {"type": ["string", "null"]},
    },
    "required": ["toetsen"],
}


# ── Tekst-extractie ──────────────────────────────────────────────────────────

def extract_text_docx(path: Path) -> str:
    doc = Document(str(path))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_text_pdf(path: Path) -> str:
    doc = fitz.open(str(path))
    parts = []
    for page in doc:
        parts.append(page.get_text())
    doc.close()
    return "\n".join(parts)


def extract_text_txt(path: Path) -> str:
    encodings = ["utf-8", "latin-1", "cp1252"]
    for enc in encodings:
        try:
            return path.read_text(encoding=enc)
        except (UnicodeDecodeError, UnicodeError):
            continue
    return path.read_text(encoding="utf-8", errors="replace")


def extract_text_xlsx(path: Path) -> str:
    wb = openpyxl.load_workbook(str(path), data_only=True)
    parts = []
    for sheet in wb.sheetnames:
        ws = wb[sheet]
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            if any(cells):
                parts.append(" | ".join(cells))
    wb.close()
    return "\n".join(parts)


def extract_text(path: Path) -> str | None:
    ext = path.suffix.lower()
    try:
        if ext == ".docx":
            return extract_text_docx(path)
        elif ext == ".pdf":
            return extract_text_pdf(path)
        elif ext == ".txt":
            return extract_text_txt(path)
        elif ext == ".xlsx":
            return extract_text_xlsx(path)
    except Exception as e:
        print(f"  FOUT bij extractie {path.name}: {e}", file=sys.stderr)
    return None


# ── Metadata-extractie uit padnaam ────────────────────────────────────────────

def parse_metadata_from_path(path: Path) -> dict:
    """Haal klas, module, vak af uit het pad van het bestand."""
    rel = path.relative_to(EXTRACTED_DIR)
    parts = rel.parts  # bijv. ('klas3', 'Studiewijzers', '3M', 'Module 1', 'bestand.pdf')

    meta = {
        "vak": None, "vak_afkorting": None, "klas": None,
        "leerjaar": None, "module": None, "docent_code": None,
        "schooljaar": "2025-2026",
    }

    # Zoek module nummer
    for p in parts:
        m = re.match(r"Module\s+(\d+)", p, re.IGNORECASE)
        if m:
            meta["module"] = int(m.group(1))
            break

    # Bepaal structuur: onderbouw vs bovenbouw
    # Onderbouw: klas*/Studiewijzers/{KLAS}/Module {N}/bestand
    # Bovenbouw: klas*/Studiewijzers/Module {N}/{VAK}/bestand
    if len(parts) >= 4 and parts[1] == "Studiewijzers":
        # Check of part[2] een klasnaam is (bijv. "1G", "3M")
        klas_match = re.match(r"^(\d+)([A-Za-z]+)$", parts[2])
        if klas_match:
            # Onderbouw patroon
            meta["klas"] = parts[2].upper()
            meta["leerjaar"] = int(klas_match.group(1))
            # Vak moet uit bestandsnaam komen
            meta["vak"] = _guess_vak_from_filename(path.stem)
        elif parts[2].startswith("Module"):
            # Bovenbouw patroon
            zip_dir = parts[0]  # "klas4" of "klas5"
            leerjaar_match = re.search(r"(\d+)", zip_dir)
            if leerjaar_match:
                meta["leerjaar"] = int(leerjaar_match.group(1))
            if len(parts) >= 5:
                candidate_vak = parts[3]
                # Guard: mapnamen die eruitzien als bestandsnamen negeren
                if not re.search(r"\.(pdf|docx|xlsx|txt|zip|doc)$", candidate_vak, re.IGNORECASE):
                    meta["vak"] = candidate_vak  # Vaknaam als directory
            # Klas uit bestandsnaam proberen
            meta["klas"] = _guess_klas_from_filename(path.stem, meta["leerjaar"])

    # Vak afkorting opzoeken
    if meta["vak"]:
        vak_lower = meta["vak"].lower()
        meta["vak_afkorting"] = VAK_AFKORTINGEN.get(vak_lower)

    # Docent code uit bestandsnaam proberen (3 hoofdletters aan einde)
    doc_match = re.search(r"\b([A-Z]{3})\b", path.stem)
    if doc_match:
        code = doc_match.group(1)
        # Filter veelvoorkomende niet-docentcodes
        if code not in {"BIO", "AKS", "ENG", "MOD", "PER", "SGL", "VWO", "GYM", "KLA"}:
            meta["docent_code"] = code

    return meta


def _guess_vak_from_filename(stem: str) -> str | None:
    """Probeer het vak af te leiden uit de bestandsnaam."""
    stem_lower = stem.lower()
    vak_keywords = {
        "wiskunde": "wiskunde", "wis ": "wiskunde", "wi ": "wiskunde",
        "nederlands": "Nederlands", "netl": "Nederlands", " ned ": "Nederlands",
        "engels": "Engels", "english": "Engels", "entl": "Engels",
        "studyplanner": "Engels", "planner m": "Engels",
        "frans": "Frans", "français": "Frans", "francais": "Frans",
        "fran": "Frans",  # voor encoding-beschadigde bestandsnamen
        "duits": "Duits", "deutsch": "Duits",
        "latijn": "Latijn", "latin": "Latijn",
        "grieks": "Grieks", "greek": "Grieks",
        "geschiedenis": "Geschiedenis", " ges ": "Geschiedenis", "gs ": "Geschiedenis",
        "aardrijkskunde": "Aardrijkskunde", " ak ": "Aardrijkskunde",
        "biologie": "Biologie", " bio ": "Biologie",
        "scheikunde": "Scheikunde", "schk": "Scheikunde",
        "natuurkunde": "Natuurkunde", " nat ": "Natuurkunde", " na ": "Natuurkunde",
        "economie": "Economie", "econ": "Economie",
        "muziek": "Muziek", "music": "Muziek",
        "drama": "Drama",
        "filosofie": "Filosofie",
        "maatschappijleer": "Maatschappijleer",
        "kunst": "Kunst",
        "beeldende": "Kunst - BV", " bv ": "Kunst - BV",
        "onderzoek": "Onderzoek",
        " gr ": "Grieks", " la ": "Latijn", "spqr": "Latijn", "argo": "Grieks",
        "bevo": "Kunst - BV", "begrippengids": "Kunst - BV",
        " nl ": "Nederlands", "netl": "Nederlands",
    }
    for keyword, vak in vak_keywords.items():
        if keyword in stem_lower or stem_lower.startswith(keyword.strip()):
            return vak
    return None


def _guess_klas_from_filename(stem: str, leerjaar: int | None) -> str | None:
    """Probeer de klas te achterhalen uit de bestandsnaam voor bovenbouw."""
    # Zoek patronen als "4econ1", "5nat2", "6WisA", "klas 4", "k5", etc.
    patterns = [
        r"(?:klas|k)\s*(\d[a-zA-Z]*)",
        r"(\d+[a-z])\s",
    ]
    for pat in patterns:
        m = re.search(pat, stem, re.IGNORECASE)
        if m:
            return m.group(1).upper()
    if leerjaar:
        return f"klas {leerjaar}"
    return None


# ── Claude API ────────────────────────────────────────────────────────────────

async def call_claude(
    client: anthropic.AsyncAnthropic,
    text: str,
    metadata: dict,
    semaphore: asyncio.Semaphore,
) -> dict:
    """Roep Claude API aan om een studiewijzer te parsen."""
    # Beperk tekst tot ~8000 chars om kosten laag te houden
    if len(text) > 8000:
        text = text[:8000] + "\n\n[... tekst ingekort ...]"

    user_msg = f"""Metadata uit bestandspad:
- Vak: {metadata.get('vak', 'onbekend')}
- Klas: {metadata.get('klas', 'onbekend')}
- Leerjaar: {metadata.get('leerjaar', 'onbekend')}
- Module: {metadata.get('module', 'onbekend')}
- Docent: {metadata.get('docent_code', 'onbekend')}

STUDIEWIJZER TEKST:
{text}"""

    async with semaphore:
        for attempt in range(3):
            try:
                response = await client.messages.create(
                    model=MODEL,
                    max_tokens=2048,
                    system=SYSTEM_PROMPT,
                    tools=[{
                        "name": "studiewijzer_output",
                        "description": "Output de geëxtraheerde toetsen uit de studiewijzer",
                        "input_schema": JSON_SCHEMA,
                    }],
                    tool_choice={"type": "tool", "name": "studiewijzer_output"},
                    messages=[{"role": "user", "content": user_msg}],
                )

                for block in response.content:
                    if block.type == "tool_use":
                        return block.input

                return {"toetsen": [], "opmerkingen": "Geen tool_use in response"}

            except anthropic.RateLimitError:
                wait = 2 ** attempt * 5
                print(f"  Rate limit, wacht {wait}s...", file=sys.stderr)
                await asyncio.sleep(wait)
            except Exception as e:
                if attempt == 2:
                    return {"toetsen": [], "opmerkingen": f"API fout: {e}"}
                await asyncio.sleep(2)

    return {"toetsen": [], "opmerkingen": "Max retries bereikt"}


# ── Hoofdlogica ───────────────────────────────────────────────────────────────

def find_studiewijzers() -> list[Path]:
    """Vind alle parseerbare studiewijzer-bestanden."""
    files = []
    for path in sorted(EXTRACTED_DIR.rglob("*")):
        if path.is_file() and path.suffix.lower() not in SKIP_EXTENSIONS:
            if path.suffix.lower() in {".docx", ".pdf", ".txt", ".xlsx"}:
                files.append(path)
    return files


def dry_run(files: list[Path], filter_pattern: str | None = None):
    """Toon alle bestanden met metadata en tekstlengte zonder API calls."""
    skipped_exts = {}
    for path in sorted(EXTRACTED_DIR.rglob("*")):
        if path.is_file() and path.suffix.lower() in SKIP_EXTENSIONS:
            skipped_exts[path.suffix.lower()] = skipped_exts.get(path.suffix.lower(), 0) + 1

    print(f"\n{'='*80}")
    print(f"STUDIEWIJZER STANDARDIZER - DRY RUN")
    print(f"{'='*80}")
    print(f"Gevonden bestanden: {len(files)}")
    print(f"Overgeslagen extensies: {skipped_exts}")
    print()

    total_chars = 0
    vak_counts = {}
    leerjaar_counts = {}

    for f in files:
        if filter_pattern and not re.search(filter_pattern, str(f)):
            continue
        meta = parse_metadata_from_path(f)
        text = extract_text(f)
        text_len = len(text) if text else 0
        total_chars += text_len

        vak = meta.get("vak") or "?"
        vak_counts[vak] = vak_counts.get(vak, 0) + 1
        lj = meta.get("leerjaar") or "?"
        leerjaar_counts[lj] = leerjaar_counts.get(lj, 0) + 1

        status = "OK" if text and text_len > 50 else "KORT" if text else "LEEG"
        print(f"  [{status:4s}] {text_len:6d} chars | "
              f"lj={meta.get('leerjaar','?')} klas={meta.get('klas','?'):5s} "
              f"mod={meta.get('module','?')} vak={vak:20s} | {f.name}")

    print(f"\n{'-'*80}")
    print(f"Totaal tekst: {total_chars:,} chars (~{total_chars//4:,} tokens)")
    print(f"\nPer leerjaar: {dict(sorted(leerjaar_counts.items()))}")
    print(f"\nPer vak:")
    for vak, count in sorted(vak_counts.items(), key=lambda x: -x[1]):
        print(f"  {vak:25s} {count:3d}")


async def process_single(path: Path):
    """Verwerk één studiewijzer en toon het resultaat."""
    meta = parse_metadata_from_path(path)
    text = extract_text(path)

    if not text:
        print(f"Kon geen tekst extraheren uit {path}", file=sys.stderr)
        return

    print(f"Bestand: {path.name}")
    print(f"Metadata: {json.dumps(meta, indent=2, ensure_ascii=False)}")
    print(f"Tekst ({len(text)} chars):")
    print(text[:2000])
    if len(text) > 2000:
        print(f"... ({len(text) - 2000} chars meer)")
    print()

    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        print("ANTHROPIC_API_KEY niet gezet - alleen tekstextractie getoond.", file=sys.stderr)
        return

    client = anthropic.AsyncAnthropic(api_key=api_key)
    sem = asyncio.Semaphore(1)
    result = await call_claude(client, text, meta, sem)

    output = {
        "bron_bestand": str(path.relative_to(EXTRACTED_DIR)),
        "metadata": meta,
        **result,
    }

    print(f"\n{'='*60}")
    print("CLAUDE OUTPUT:")
    print(json.dumps(output, indent=2, ensure_ascii=False))


async def process_all(filter_pattern: str | None = None, skip_existing: bool = False):
    """Verwerk alle studiewijzers."""
    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        print("ANTHROPIC_API_KEY niet gezet!", file=sys.stderr)
        sys.exit(1)

    files = find_studiewijzers()
    if filter_pattern:
        files = [f for f in files if re.search(filter_pattern, str(f))]

    if skip_existing:
        before = len(files)
        files = [f for f in files
                 if not (OUTPUT_DIR / f.relative_to(EXTRACTED_DIR).with_suffix(".json")).exists()]
        print(f"Skip-existing: {before - len(files)} al verwerkt, {len(files)} nieuw")

    print(f"Verwerken van {len(files)} studiewijzers...")
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    client = anthropic.AsyncAnthropic(api_key=api_key)
    semaphore = asyncio.Semaphore(MAX_CONCURRENT)

    summary = {
        "totaal_bestanden": len(files),
        "verwerkt": 0,
        "fouten": 0,
        "totaal_toetsen": 0,
        "per_leerjaar": {},
        "per_vak": {},
        "per_type": {},
    }

    start_time = time.time()

    async def process_one(path: Path, idx: int):
        meta = parse_metadata_from_path(path)
        text = extract_text(path)

        if not text or len(text) < 20:
            print(f"  [{idx+1}/{len(files)}] SKIP (geen tekst): {path.name}")
            summary["fouten"] += 1
            return

        result = await call_claude(client, text, meta, semaphore)

        output = {
            "bron_bestand": str(path.relative_to(EXTRACTED_DIR)),
            "metadata": meta,
            **result,
        }

        # Bewaar JSON
        rel = path.relative_to(EXTRACTED_DIR)
        out_path = OUTPUT_DIR / rel.with_suffix(".json")
        out_path.parent.mkdir(parents=True, exist_ok=True)
        out_path.write_text(json.dumps(output, indent=2, ensure_ascii=False), encoding="utf-8")

        n_toetsen = len(result.get("toetsen", []))
        summary["verwerkt"] += 1
        summary["totaal_toetsen"] += n_toetsen

        lj = str(meta.get("leerjaar", "?"))
        summary["per_leerjaar"][lj] = summary["per_leerjaar"].get(lj, 0) + n_toetsen

        vak = meta.get("vak") or "onbekend"
        summary["per_vak"][vak] = summary["per_vak"].get(vak, 0) + n_toetsen

        for toets in result.get("toetsen", []):
            if not isinstance(toets, dict):
                continue
            t = toets.get("type", "anders")
            summary["per_type"][t] = summary["per_type"].get(t, 0) + 1

        print(f"  [{idx+1}/{len(files)}] {n_toetsen} toetsen: {path.name}")

    # Verwerk in batches
    tasks = [process_one(f, i) for i, f in enumerate(files)]
    await asyncio.gather(*tasks)

    elapsed = time.time() - start_time
    summary["duur_seconden"] = round(elapsed, 1)

    # Schrijf samenvatting
    summary_path = OUTPUT_DIR / "summary.json"
    summary_path.write_text(json.dumps(summary, indent=2, ensure_ascii=False), encoding="utf-8")

    print(f"\n{'='*60}")
    print(f"KLAAR in {elapsed:.1f}s")
    print(f"Verwerkt: {summary['verwerkt']}/{summary['totaal_bestanden']}")
    print(f"Fouten: {summary['fouten']}")
    print(f"Totaal toetsen gevonden: {summary['totaal_toetsen']}")
    print(f"\nPer type: {json.dumps(summary['per_type'], indent=2)}")
    print(f"\nOutput: {OUTPUT_DIR}")
    print(f"Samenvatting: {summary_path}")


def main():
    parser = argparse.ArgumentParser(description="Studiewijzer Standardizer")
    parser.add_argument("--dry-run", action="store_true",
                        help="Toon bestanden en metadata zonder API calls")
    parser.add_argument("--single", type=str,
                        help="Verwerk één bestand (pad relatief t.o.v. _extracted/)")
    parser.add_argument("--filter", type=str,
                        help="Regex filter op bestandspaden (bijv. 'klas3.*Module 1')")
    parser.add_argument("--skip-existing", action="store_true",
                        help="Sla bestanden over die al een output JSON hebben")
    args = parser.parse_args()

    files = find_studiewijzers()

    if args.dry_run:
        dry_run(files, args.filter)
    elif args.single:
        single_path = EXTRACTED_DIR / args.single
        if not single_path.exists():
            # Probeer als absoluut pad
            single_path = Path(args.single)
        if not single_path.exists():
            print(f"Bestand niet gevonden: {args.single}", file=sys.stderr)
            sys.exit(1)
        asyncio.run(process_single(single_path))
    else:
        asyncio.run(process_all(args.filter, skip_existing=args.skip_existing))


if __name__ == "__main__":
    main()
